#!/usr/bin/env python3
"""Promote approved reader prose into a canonical DOCX.

The illustrated HTML pages remain publishing derivatives. This tool exists so an
approved editorial diff can be applied to a canonical manuscript without
rebuilding or replacing untargeted chapters.

Requires: python-docx
"""

from __future__ import annotations

import argparse
import re
import tempfile
from copy import deepcopy
from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ONES = {
    "ONE": 1,
    "TWO": 2,
    "THREE": 3,
    "FOUR": 4,
    "FIVE": 5,
    "SIX": 6,
    "SEVEN": 7,
    "EIGHT": 8,
    "NINE": 9,
    "TEN": 10,
    "ELEVEN": 11,
    "TWELVE": 12,
    "THIRTEEN": 13,
    "FOURTEEN": 14,
    "FIFTEEN": 15,
    "SIXTEEN": 16,
    "SEVENTEEN": 17,
    "EIGHTEEN": 18,
    "NINETEEN": 19,
}
TENS = {
    "TWENTY": 20,
    "THIRTY": 30,
    "FORTY": 40,
    "FIFTY": 50,
    "SIXTY": 60,
    "SEVENTY": 70,
    "EIGHTY": 80,
}
HEADING_RE = re.compile(r"^CHAPTER\s+([^\n]+)(?:\n.*)?$", re.IGNORECASE)


class ProseParagraphParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.article_depth = 0
        self.in_prose = False
        self.in_p = False
        self.current: list[str] = []
        self.paragraphs: list[str] = []

    def handle_starttag(self, tag: str, attrs):
        attrs_dict = dict(attrs)
        classes = set(attrs_dict.get("class", "").split())
        if tag == "article" and "prose" in classes and not self.in_prose:
            self.in_prose = True
            self.article_depth = 1
            return
        if self.in_prose:
            self.article_depth += 1
            if tag == "p" and not self.in_p:
                self.in_p = True
                self.current = []
            elif tag == "br" and self.in_p:
                self.current.append("\n")

    def handle_endtag(self, tag: str):
        if not self.in_prose:
            return
        if tag == "p" and self.in_p:
            self.paragraphs.append("".join(self.current))
            self.current = []
            self.in_p = False
        self.article_depth -= 1
        if self.article_depth == 0:
            self.in_prose = False

    def handle_data(self, data: str):
        if self.in_prose and self.in_p:
            self.current.append(data)


def _number_from_heading(text: str) -> int | None:
    match = HEADING_RE.match(text.strip())
    if not match:
        return None
    token = match.group(1).strip().upper()
    if token.isdigit():
        return int(token)
    if token in ONES:
        return ONES[token]
    parts = token.replace("-", " ").split()
    if len(parts) == 1 and parts[0] in TENS:
        return TENS[parts[0]]
    if len(parts) == 2 and parts[0] in TENS and parts[1] in ONES:
        return TENS[parts[0]] + ONES[parts[1]]
    return None


def _heading_map(doc: Document) -> dict[int, int]:
    result: dict[int, int] = {}
    for index, paragraph in enumerate(doc.paragraphs):
        number = _number_from_heading(paragraph.text)
        if number is not None:
            if number in result:
                raise ValueError(f"duplicate chapter heading for {number}")
            result[number] = index
    return result


def _chapter_bodies(doc: Document) -> dict[int, list[str]]:
    headings = _heading_map(doc)
    ordered = sorted(headings.items(), key=lambda item: item[1])
    paragraphs = doc.paragraphs
    bodies: dict[int, list[str]] = {}
    for pos, (chapter, start) in enumerate(ordered):
        end = ordered[pos + 1][1] if pos + 1 < len(ordered) else len(paragraphs)
        bodies[chapter] = [p.text for p in paragraphs[start + 1 : end]]
    return bodies


def extract_html_prose(path: Path) -> list[str]:
    parser = ProseParagraphParser()
    parser.feed(path.read_text(encoding="utf-8"))
    if not parser.paragraphs:
        raise ValueError(f"no article.prose paragraphs found in {path}")
    return parser.paragraphs


def _make_paragraph(text: str, template_element):
    p = OxmlElement("w:p")
    p_pr = template_element.find(qn("w:pPr"))
    if p_pr is not None:
        p.append(deepcopy(p_pr))
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    if text.startswith(" ") or text.endswith(" "):
        t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    p.append(run)
    return p


def promote_chapters(docx_path: Path, chapters_dir: Path, chapters: list[int]) -> bool:
    docx_path = Path(docx_path)
    chapters_dir = Path(chapters_dir)
    requested = sorted(set(chapters))
    if not requested:
        raise ValueError("no chapters requested")

    replacements = {
        chapter: extract_html_prose(chapters_dir / f"{chapter:03d}.html")
        for chapter in requested
    }
    for chapter, paragraphs in replacements.items():
        if any("—" in paragraph for paragraph in paragraphs):
            raise ValueError(f"chapter {chapter} contains an em dash")

    doc = Document(docx_path)
    before = _chapter_bodies(doc)
    heading_count = len(_heading_map(doc))
    for chapter in requested:
        if chapter not in before:
            raise ValueError(f"chapter {chapter} not found in canonical DOCX")

    if all(before[chapter] == replacements[chapter] for chapter in requested):
        return False

    for chapter in sorted(requested, reverse=True):
        headings = _heading_map(doc)
        start = headings[chapter]
        later = sorted(index for index in headings.values() if index > start)
        end = later[0] if later else len(doc.paragraphs)
        paragraphs = doc.paragraphs
        heading = paragraphs[start]
        body = paragraphs[start + 1 : end]
        template = body[0]._p if body else heading._p

        for paragraph in body:
            paragraph._p.getparent().remove(paragraph._p)

        if later:
            # Interior chapter: preserve the following chapter heading as the
            # insertion boundary.
            next_heading = doc.paragraphs[_heading_map(doc)[chapter] + 1]
            # The expression above is not guaranteed to identify the next
            # chapter after body removal, so resolve it from the next chapter
            # number instead.
            current_headings = _heading_map(doc)
            current_start = current_headings[chapter]
            next_indices = sorted(index for index in current_headings.values() if index > current_start)
            next_heading = doc.paragraphs[next_indices[0]]
            for text in replacements[chapter]:
                next_heading._p.addprevious(_make_paragraph(text, template))
        else:
            # Final chapter: there is no following heading to anchor insertion.
            # Insert each new paragraph immediately after the chapter heading in
            # reverse order. This keeps Word's trailing section properties in
            # place and preserves the requested paragraph order.
            for text in reversed(replacements[chapter]):
                heading._p.addnext(_make_paragraph(text, template))

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False, dir=docx_path.parent) as tmp:
        tmp_path = Path(tmp.name)
    try:
        doc.save(tmp_path)
        check = Document(tmp_path)
        after = _chapter_bodies(check)
        if len(_heading_map(check)) != heading_count:
            raise ValueError("chapter heading count changed during promotion")
        for chapter in requested:
            if after[chapter] != replacements[chapter]:
                raise ValueError(f"chapter {chapter} did not round-trip to reader prose")
        for chapter, paragraphs in before.items():
            if chapter not in requested and after.get(chapter) != paragraphs:
                raise ValueError(f"untargeted chapter {chapter} changed")
        tmp_path.replace(docx_path)
        return True
    finally:
        if tmp_path.exists():
            tmp_path.unlink()


def _parse_range(value: str) -> list[int]:
    if "-" in value:
        start, end = (int(part) for part in value.split("-", 1))
        if end < start:
            raise argparse.ArgumentTypeError("chapter range must be ascending")
        return list(range(start, end + 1))
    return [int(value)]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("range", help="chapter number or inclusive range, e.g. 1-5")
    parser.add_argument(
        "--docx",
        type=Path,
        default=Path("state/manuscript/Peg_Leg_Greg_authoritative_ch82_final_name_map.docx"),
    )
    parser.add_argument("--chapters-dir", type=Path, default=Path("chapters"))
    args = parser.parse_args()
    chapters = _parse_range(args.range)
    changed = promote_chapters(args.docx, args.chapters_dir, chapters)
    if changed:
        print(f"promoted chapters {chapters[0]}-{chapters[-1]} into {args.docx}")
    else:
        print(f"chapters {chapters[0]}-{chapters[-1]} already synchronized in {args.docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
