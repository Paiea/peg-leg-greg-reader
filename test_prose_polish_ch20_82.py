import hashlib
import re
import unittest
from html.parser import HTMLParser
from pathlib import Path

from docx import Document


ROOT = Path(__file__).parent
BOOK1 = ROOT / "state/manuscript/Peg_Leg_Greg_authoritative_ch82_final_name_map.docx"
ORIGINAL_BOOK1_SHA = "cbccfde5bf54ab7afd427505ef5f48d3a9c5f7f50552e69e9616b4294b7339d8"


class ProseParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.in_prose = False
        self.paragraphs = 0

    def handle_starttag(self, tag, attrs):
        if tag == "article" and "prose" in dict(attrs).get("class", ""):
            self.in_prose = True
        elif self.in_prose and tag == "p":
            self.paragraphs += 1

    def handle_endtag(self, tag):
        if tag == "article":
            self.in_prose = False


class ProsePolishTest(unittest.TestCase):
    def test_polished_book1_is_readable_and_complete(self):
        self.assertNotEqual(hashlib.sha256(BOOK1.read_bytes()).hexdigest(), ORIGINAL_BOOK1_SHA)
        document = Document(BOOK1)
        headings = [p.text for p in document.paragraphs if p.text.startswith("CHAPTER ") and p.text == p.text.upper()]
        self.assertEqual(len(headings), 82)
        self.assertIn("THE BOY", headings[0])
        self.assertIn("THE RECONCILER", headings[-1])

    def test_reader_range_has_recomposed_rhythm_without_losing_chapters(self):
        total_paragraphs = 0
        for number in range(20, 83):
            page = ROOT / "chapters" / f"{number:03}.html"
            parser = ProseParser()
            parser.feed(page.read_text(encoding="utf-8"))
            self.assertGreater(parser.paragraphs, 0, page)
            total_paragraphs += parser.paragraphs
        self.assertLess(total_paragraphs, 42000)
        self.assertEqual(len(list((ROOT / "chapters").glob("[0-9][0-9][0-9].html"))), 137)


if __name__ == "__main__":
    unittest.main()
