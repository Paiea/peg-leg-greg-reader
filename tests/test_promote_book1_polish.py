import tempfile
import time
import unittest
from pathlib import Path

try:
    from docx import Document
    from scripts.promote_book1_polish import _number_from_heading, promote_chapters
except ModuleNotFoundError:
    Document = None
    _number_from_heading = None
    promote_chapters = None


@unittest.skipIf(Document is None, "python-docx is only required for manuscript source promotion")
class PromoteBook1PolishTests(unittest.TestCase):
    def make_docx(self, path: Path) -> None:
        doc = Document()
        doc.add_paragraph("PEG-LEG GREG")
        doc.add_paragraph("CHAPTER ONE\nTHE BOY")
        doc.add_paragraph("Old one.")
        doc.add_paragraph("Old two.")
        doc.add_paragraph("CHAPTER TWO\nTHE BORROWER")
        doc.add_paragraph("Keep this chapter exactly.")
        doc.add_paragraph("CHAPTER THREE\nTHE INVESTOR")
        doc.add_paragraph("Keep this too.")
        doc.save(path)

    def test_promotes_only_requested_chapter(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx = root / "book1.docx"
            chapters = root / "chapters"
            chapters.mkdir()
            self.make_docx(docx)
            (chapters / "001.html").write_text(
                '<article class="prose"><p>New one.</p><p>New two.</p><p>New three.</p></article>',
                encoding="utf-8",
            )
            changed = promote_chapters(docx, chapters, [1])
            self.assertTrue(changed)
            reopened = Document(docx)
            text = [p.text for p in reopened.paragraphs]
            self.assertEqual(text, ["PEG-LEG GREG", "CHAPTER ONE\nTHE BOY", "New one.", "New two.", "New three.", "CHAPTER TWO\nTHE BORROWER", "Keep this chapter exactly.", "CHAPTER THREE\nTHE INVESTOR", "Keep this too."])

    def test_promotes_final_chapter_without_following_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx = root / "book1.docx"
            chapters = root / "chapters"
            chapters.mkdir()
            self.make_docx(docx)
            (chapters / "003.html").write_text('<article class="prose"><p>New final one.</p><p>New final two.</p></article>', encoding="utf-8")
            changed = promote_chapters(docx, chapters, [3])
            self.assertTrue(changed)
            reopened = Document(docx)
            text = [p.text for p in reopened.paragraphs]
            self.assertEqual(text, ["PEG-LEG GREG", "CHAPTER ONE\nTHE BOY", "Old one.", "Old two.", "CHAPTER TWO\nTHE BORROWER", "Keep this chapter exactly.", "CHAPTER THREE\nTHE INVESTOR", "New final one.", "New final two."])

    def test_recognizes_ninety_series_heading(self):
        self.assertEqual(_number_from_heading("CHAPTER NINETY\nTHE EXAMPLE"), 90)
        self.assertEqual(_number_from_heading("CHAPTER NINETY-NINE\nTHE EXAMPLE"), 99)

    def test_recognizes_hundred_series_heading(self):
        self.assertEqual(_number_from_heading("CHAPTER ONE HUNDRED\nTHE EXAMPLE"), 100)
        self.assertEqual(_number_from_heading("CHAPTER ONE HUNDRED ONE\nTHE EXAMPLE"), 101)
        self.assertEqual(_number_from_heading("CHAPTER ONE HUNDRED THIRTY-SEVEN\nTHE EXAMPLE"), 137)

    def test_rejects_em_dash_in_promoted_prose(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx = root / "book1.docx"
            chapters = root / "chapters"
            chapters.mkdir()
            self.make_docx(docx)
            (chapters / "001.html").write_text('<article class="prose"><p>Bad — dash.</p></article>', encoding="utf-8")
            with self.assertRaisesRegex(ValueError, "em dash"):
                promote_chapters(docx, chapters, [1])

    def test_already_synchronized_chapter_does_not_touch_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            docx = root / "book1.docx"
            chapters = root / "chapters"
            chapters.mkdir()
            self.make_docx(docx)
            (chapters / "001.html").write_text('<article class="prose"><p>New one.</p><p>New two.</p></article>', encoding="utf-8")
            self.assertTrue(promote_chapters(docx, chapters, [1]))
            before = docx.stat().st_mtime_ns
            time.sleep(0.02)
            changed = promote_chapters(docx, chapters, [1])
            after = docx.stat().st_mtime_ns
            self.assertFalse(changed)
            self.assertEqual(before, after)


if __name__ == "__main__":
    unittest.main()
